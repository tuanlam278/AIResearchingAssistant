import io
import logging
from pathlib import Path
from typing import Any, Dict, List


from app.services.document_structure_service import DocumentBlock, blocks_to_page, classify_markdown_block, normalize_plain_text, structure_flat_pages
from app.services.pdf_table_extractor import table_to_markdown
from app.services.pdf_parser import parse_pdf

logger = logging.getLogger(__name__)

SUPPORTED_TEXT_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}
ACCEPTED_UPLOAD_EXTENSIONS = SUPPORTED_TEXT_EXTENSIONS | {".doc", ".rtf"}


class UnsupportedDocumentType(ValueError):
    """Raised when the uploaded file type is intentionally not indexed."""


class EmptyDocumentText(ValueError):
    """Raised when parsing succeeds but no useful text is extracted."""


def get_file_extension(filename: str | None) -> str:
    return Path(filename or "").suffix.lower()


def get_file_type(filename: str | None) -> str:
    ext = get_file_extension(filename)
    return ext.lstrip(".") or "unknown"


def validate_research_file(filename: str | None) -> str:
    ext = get_file_extension(filename)
    if ext not in ACCEPTED_UPLOAD_EXTENSIONS:
        raise UnsupportedDocumentType("Định dạng file chưa được hỗ trợ. Vui lòng upload PDF, DOCX, TXT hoặc MD.")
    if ext == ".doc":
        raise UnsupportedDocumentType("File .doc cũ chưa được hỗ trợ. Vui lòng chuyển sang .docx hoặc PDF.")
    if ext == ".rtf":
        raise UnsupportedDocumentType("File .rtf chưa được hỗ trợ. Vui lòng chuyển sang DOCX, PDF, TXT hoặc MD.")
    return ext


def _ensure_non_empty(pages: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    text = "\n".join(str(page.get("content") or "").strip() for page in pages).strip()
    if not text:
        raise EmptyDocumentText("Không đọc được nội dung văn bản từ file này.")
    return pages


def _decode_text(contents: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "cp1258", "latin-1"):
        try:
            return contents.decode(encoding)
        except UnicodeDecodeError:
            continue
    return contents.decode("utf-8", errors="replace")


def _parse_text(contents: bytes) -> List[Dict[str, Any]]:
    text = _decode_text(contents).replace("\r\n", "\n").replace("\r", "\n").strip()
    return _ensure_non_empty(structure_flat_pages([{"page_number": 1, "content": text, "source": "text"}], source="text"))


def _parse_docx(contents: bytes) -> List[Dict[str, Any]]:
    try:
        from docx import Document as DocxDocument
    except ImportError as exc:
        raise UnsupportedDocumentType("Máy chủ chưa cài python-docx để đọc DOCX.") from exc
    document = DocxDocument(io.BytesIO(contents))
    doc_blocks: list[DocumentBlock] = []
    current_section: str | None = None
    for paragraph in document.paragraphs:
        value = paragraph.text.strip()
        if value:
            block_type = classify_markdown_block(value)
            if block_type == "heading":
                current_section = value.lstrip("#").strip()[:180]
            doc_blocks.append(DocumentBlock(page=1, block_index=len(doc_blocks), block_type=block_type, section=current_section, markdown=value, text=normalize_plain_text(value), bbox=None, confidence=None, source="docx"))
    try:
        for table in document.tables:
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            markdown = table_to_markdown(rows)
            if markdown:
                doc_blocks.append(DocumentBlock(page=1, block_index=len(doc_blocks), block_type="table", section=current_section, markdown=markdown, text=normalize_plain_text(markdown), bbox=None, confidence=None, source="docx"))
    except Exception as exc:
        logger.warning("DOCX table extraction failed; continuing with paragraphs only: %s", exc)
    return _ensure_non_empty([blocks_to_page(1, doc_blocks)])


async def parse_document(contents: bytes, filename: str | None) -> tuple[List[Dict[str, Any]], str]:
    """Parse a supported research document into chunker-compatible page dictionaries."""
    ext = validate_research_file(filename)
    if ext == ".pdf":
        return _ensure_non_empty(await parse_pdf(contents)), "pdf"
    if ext == ".docx":
        return _parse_docx(contents), "docx"
    if ext in {".txt", ".md"}:
        return _parse_text(contents), ext.lstrip(".")
    raise UnsupportedDocumentType("Định dạng file chưa được hỗ trợ. Vui lòng upload PDF, DOCX, TXT hoặc MD.")

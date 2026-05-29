import logging
import re
import unicodedata
from datetime import datetime, timezone
from typing import Any, List
from urllib.parse import quote

from fastapi import APIRouter, Depends, HTTPException, status
from pydantic import BaseModel, Field

from app.db.supabase_client import supabase
from app.dependencies import get_current_user
from app.services.embedder import embed_query
from app.services.retriever import OUT_OF_SCOPE_WARNING, retrieve_rag_context

logger = logging.getLogger(__name__)
router = APIRouter(tags=["research-sessions"])


class CreateResearchSessionRequest(BaseModel):
    selected_document_ids: List[str] = Field(..., min_length=1, max_length=50)


class UpdateResearchSessionRequest(BaseModel):
    title: str | None = Field(default=None, min_length=1, max_length=200)
    is_starred: bool | None = None


def _supabase_response_data(resp: Any) -> tuple[Any, Any]:
    if isinstance(resp, dict):
        return resp.get("data"), resp.get("error")
    return getattr(resp, "data", None), getattr(resp, "error", None)


def _get_user_id(user: dict) -> str:
    user_id = user.get("id") or user.get("user_id")
    if not user_id:
        raise HTTPException(status_code=status.HTTP_401_UNAUTHORIZED, detail={"code": "UNAUTHORIZED", "message": "Token không hợp lệ"})
    return user_id


def _ensure_notebook_owner(notebook_id: str, user_id: str) -> None:
    try:
        resp = supabase.table("notebooks").select("id").match({"id": notebook_id, "user_id": user_id}).execute()
    except Exception as exc:
        logger.exception("Notebook ownership check failed")
        raise HTTPException(status_code=500, detail={"code": "INTERNAL_ERROR", "message": "Lỗi khi kiểm tra notebook"}) from exc
    data, error = _supabase_response_data(resp)
    if error:
        raise HTTPException(status_code=500, detail={"code": "INTERNAL_ERROR", "message": "Lỗi khi kiểm tra notebook"})
    if not data:
        raise HTTPException(status_code=404, detail={"code": "DOC_NOT_FOUND", "message": "Không tìm thấy notebook"})


def _load_selected_documents(notebook_id: str, document_ids: list[str]) -> list[dict]:
    unique_ids = list(dict.fromkeys(str(doc_id) for doc_id in document_ids if doc_id))
    if not unique_ids:
        raise HTTPException(status_code=400, detail={"code": "NO_DOCUMENT_SELECTED", "message": "Vui lòng chọn ít nhất một tài liệu để nghiên cứu."})
    try:
        resp = (
            supabase.table("documents")
            .select("id, filename, page_count, chunk_count, created_at")
            .eq("notebook_id", notebook_id)
            .in_("id", unique_ids)
            .execute()
        )
    except Exception as exc:
        logger.exception("Could not load selected documents")
        raise HTTPException(status_code=500, detail={"code": "INTERNAL_ERROR", "message": "Lỗi khi kiểm tra tài liệu"}) from exc
    rows, error = _supabase_response_data(resp)
    if error:
        raise HTTPException(status_code=500, detail={"code": "INTERNAL_ERROR", "message": "Lỗi khi kiểm tra tài liệu"})
    if len(rows or []) != len(unique_ids):
        raise HTTPException(status_code=400, detail={"code": "INVALID_SELECTED_DOCUMENTS", "message": "Một hoặc nhiều tài liệu đã bị xóa hoặc không còn hợp lệ."})
    order = {doc_id: index for index, doc_id in enumerate(unique_ids)}
    return sorted(rows or [], key=lambda row: order.get(str(row["id"]), 0))


def _normalize_session(row: dict) -> dict:
    return {
        "id": row.get("id"),
        "notebook_id": row.get("notebook_id"),
        "title": row.get("title"),
        "selected_document_ids": row.get("selected_document_ids") or [],
        "created_at": row.get("created_at"),
        "updated_at": row.get("updated_at"),
        "is_starred": row.get("is_starred", False),
    }


@router.get("/workspaces/{workspace_id}/research-sessions", response_model=dict)
async def list_research_sessions(workspace_id: str, user: dict = Depends(get_current_user)):
    user_id = _get_user_id(user)
    _ensure_notebook_owner(workspace_id, user_id)
    try:
        resp = (
            supabase.table("research_sessions")
            .select("id, notebook_id, title, selected_document_ids, created_at, updated_at, is_starred")
            .eq("notebook_id", workspace_id)
            .order("is_starred", desc=True)
            .order("created_at", desc=True)
            .execute()
        )
    except Exception as exc:
        logger.exception("List research sessions failed")
        raise HTTPException(status_code=500, detail={"code": "INTERNAL_ERROR", "message": "Lỗi khi lấy lịch sử nghiên cứu"}) from exc
    rows, error = _supabase_response_data(resp)
    if error:
        raise HTTPException(status_code=500, detail={"code": "INTERNAL_ERROR", "message": "Lỗi khi lấy lịch sử nghiên cứu"})
    sessions = [_normalize_session(row) for row in rows or []]
    return {"success": True, "data": {"sessions": sessions, "total": len(sessions)}}


@router.post("/workspaces/{workspace_id}/research-sessions", response_model=dict)
async def create_research_session(workspace_id: str, body: CreateResearchSessionRequest, user: dict = Depends(get_current_user)):
    user_id = _get_user_id(user)
    _ensure_notebook_owner(workspace_id, user_id)
    selected_docs = _load_selected_documents(workspace_id, body.selected_document_ids)
    selected_ids = [str(doc["id"]) for doc in selected_docs]
    filenames = [doc.get("filename") or "Tài liệu" for doc in selected_docs]

    try:
        count_resp = supabase.table("research_sessions").select("id").eq("notebook_id", workspace_id).execute()
    except Exception as exc:
        logger.exception("Research session count failed")
        raise HTTPException(status_code=500, detail={"code": "INTERNAL_ERROR", "message": "Lỗi khi tạo phiên nghiên cứu"}) from exc
    existing, error = _supabase_response_data(count_resp)
    if error:
        raise HTTPException(status_code=500, detail={"code": "INTERNAL_ERROR", "message": "Lỗi khi tạo phiên nghiên cứu"})

    title = f"Nghiên cứu từ {', '.join(filenames)} - lần {len(existing or []) + 1}"
    try:
        resp = supabase.table("research_sessions").insert({
            "notebook_id": workspace_id,
            "title": title,
            "selected_document_ids": selected_ids,
            "is_starred": False,
        }).execute()
    except Exception as exc:
        logger.exception("Create research session failed")
        raise HTTPException(status_code=500, detail={"code": "INTERNAL_ERROR", "message": "Lỗi khi tạo phiên nghiên cứu"}) from exc
    rows, insert_error = _supabase_response_data(resp)
    if insert_error or not rows:
        raise HTTPException(status_code=500, detail={"code": "INTERNAL_ERROR", "message": "Lỗi khi tạo phiên nghiên cứu"})
    return {"success": True, "data": {"session": _normalize_session(rows[0])}}


@router.patch("/research-sessions/{session_id}", response_model=dict)
async def update_research_session(
    session_id: str,
    body: UpdateResearchSessionRequest,
    user: dict = Depends(get_current_user),
):
    user_id = _get_user_id(user)
    _get_owned_session(session_id, user_id)

    updates: dict[str, Any] = {"updated_at": datetime.now(timezone.utc).isoformat()}
    if body.title is not None:
        updates["title"] = body.title.strip()
    if body.is_starred is not None:
        updates["is_starred"] = body.is_starred

    if len(updates) == 1:
        raise HTTPException(status_code=400, detail={"code": "NO_UPDATES", "message": "Không có dữ liệu cập nhật"})

    try:
        resp = supabase.table("research_sessions").update(updates).eq("id", session_id).execute()
    except Exception as exc:
        logger.exception("Update research session failed")
        raise HTTPException(status_code=500, detail={"code": "INTERNAL_ERROR", "message": "Lỗi khi cập nhật lịch sử nghiên cứu"}) from exc
    rows, error = _supabase_response_data(resp)
    if error or not rows:
        raise HTTPException(status_code=500, detail={"code": "INTERNAL_ERROR", "message": "Lỗi khi cập nhật lịch sử nghiên cứu"})
    return {"success": True, "data": {"session": _normalize_session(rows[0])}}


@router.delete("/research-sessions/{session_id}", response_model=dict)
async def delete_research_session(session_id: str, user: dict = Depends(get_current_user)):
    user_id = _get_user_id(user)
    _get_owned_session(session_id, user_id)
    try:
        supabase.table("research_session_messages").delete().eq("research_session_id", session_id).execute()
        resp = supabase.table("research_sessions").delete().eq("id", session_id).execute()
    except Exception as exc:
        logger.exception("Delete research session failed")
        raise HTTPException(status_code=500, detail={"code": "INTERNAL_ERROR", "message": "Lỗi khi xoá lịch sử nghiên cứu"}) from exc
    rows, error = _supabase_response_data(resp)
    if error:
        raise HTTPException(status_code=500, detail={"code": "INTERNAL_ERROR", "message": "Lỗi khi xoá lịch sử nghiên cứu"})
    if not rows:
        raise HTTPException(status_code=404, detail={"code": "DOC_NOT_FOUND", "message": "Không tìm thấy phiên nghiên cứu"})
    return {"success": True, "data": {"session_id": session_id, "deleted": True}}


@router.get("/research-sessions/{session_id}/messages", response_model=dict)
async def get_research_session_messages(session_id: str, user: dict = Depends(get_current_user)):
    user_id = _get_user_id(user)
    session = _get_owned_session(session_id, user_id)
    try:
        resp = (
            supabase.table("research_session_messages")
            .select("id, role, content, citations, created_at")
            .eq("research_session_id", session_id)
            .order("created_at", desc=False)
            .execute()
        )
    except Exception as exc:
        logger.exception("List research session messages failed")
        raise HTTPException(status_code=500, detail={"code": "INTERNAL_ERROR", "message": "Lỗi khi lấy lịch sử chat"}) from exc
    rows, error = _supabase_response_data(resp)
    if error:
        raise HTTPException(status_code=500, detail={"code": "INTERNAL_ERROR", "message": "Lỗi khi lấy lịch sử chat"})
    messages = [
        {
            "id": row.get("id"),
            "role": row.get("role"),
            "content": row.get("content") or "",
            "citations": row.get("citations") or [],
            "created_at": row.get("created_at"),
        }
        for row in rows or []
    ]
    return {"success": True, "data": {"session": session, "messages": messages}}


@router.delete("/research-sessions/{session_id}/messages", response_model=dict)
async def clear_research_session_messages(session_id: str, user: dict = Depends(get_current_user)):
    user_id = _get_user_id(user)
    _get_owned_session(session_id, user_id)
    try:
        resp = supabase.table("research_session_messages").delete().eq("research_session_id", session_id).execute()
    except Exception as exc:
        logger.exception("Clear research session messages failed")
        raise HTTPException(status_code=500, detail={"code": "INTERNAL_ERROR", "message": "Lỗi khi xóa lịch sử cuộc trò chuyện"}) from exc
    _, error = _supabase_response_data(resp)
    if error:
        raise HTTPException(status_code=500, detail={"code": "INTERNAL_ERROR", "message": "Lỗi khi xóa lịch sử cuộc trò chuyện"})
    return {"success": True, "data": {"message": "Đã xóa lịch sử cuộc trò chuyện."}}


def _get_owned_session(session_id: str, user_id: str) -> dict:
    try:
        resp = (
            supabase.table("research_sessions")
            .select("id, notebook_id, title, selected_document_ids, created_at, updated_at, is_starred, notebooks!inner(user_id)")
            .eq("id", session_id)
            .eq("notebooks.user_id", user_id)
            .execute()
        )
    except Exception as exc:
        logger.exception("Research session ownership check failed")
        raise HTTPException(status_code=500, detail={"code": "INTERNAL_ERROR", "message": "Lỗi khi kiểm tra phiên nghiên cứu"}) from exc
    rows, error = _supabase_response_data(resp)
    if error:
        raise HTTPException(status_code=500, detail={"code": "INTERNAL_ERROR", "message": "Lỗi khi kiểm tra phiên nghiên cứu"})
    if not rows:
        raise HTTPException(status_code=404, detail={"code": "DOC_NOT_FOUND", "message": "Không tìm thấy phiên nghiên cứu"})
    return _normalize_session(rows[0])


class GenerateFlashcardsRequest(BaseModel):
    selected_document_ids: List[str] = Field(default_factory=list, max_length=50)
    count: int = Field(default=5, ge=1, le=5)


class GenerateQuizRequest(BaseModel):
    selected_document_ids: List[str] = Field(default_factory=list, max_length=50)
    count: int = Field(default=3, ge=1, le=5)
    question_type: str = Field(default="mixed")


class GenerateTestRequest(BaseModel):
    selected_document_ids: List[str] = Field(default_factory=list, min_length=1, max_length=50)
    count: int = Field(default=10)


def _selected_session_documents(session: dict, requested_ids: list[str]) -> list[str]:
    session_doc_ids = [str(doc_id) for doc_id in (session.get("selected_document_ids") or [])]
    requested_doc_ids = [str(doc_id) for doc_id in (requested_ids or session_doc_ids)]
    allowed = set(session_doc_ids)
    selected = [doc_id for doc_id in requested_doc_ids if doc_id in allowed]
    if not selected:
        raise HTTPException(status_code=400, detail={"code": "INVALID_SELECTED_DOCUMENTS", "message": "Vui lòng chọn tài liệu trước khi tạo quiz/test."})
    return selected


def _context_from_chunks(rows: list[dict]) -> str:
    return "\n\n".join(
        f"[Trang {row.get('page_number') or '?'} - {row.get('section') or 'Unknown'}] {row.get('content') or ''}"
        for row in rows
        if row.get("content")
    )


async def _load_generation_context(session: dict, selected: list[str], rag_query: str) -> tuple[str, str | None]:
    try:
        query_vector = await embed_query(rag_query)
        retrieval = await retrieve_rag_context(query_vector, session.get("notebook_id"), selected)
    except Exception as exc:
        logger.exception("Load quiz/test RAG context failed")
        raise HTTPException(status_code=500, detail={"code": "INTERNAL_ERROR", "message": "Không thể tải nội dung tài liệu."}) from exc
    rows = retrieval.chunks
    if not rows:
        raise HTTPException(status_code=404, detail={"code": "DOC_NOT_FOUND", "message": "Không tìm thấy nội dung tài liệu để tạo quiz/test."})
    return _context_from_chunks(rows), OUT_OF_SCOPE_WARNING if retrieval.is_out_of_scope else None


def _ascii_export_filename(title: str | None) -> str:
    raw = (title or f"research-chat-{datetime.now(timezone.utc).date().isoformat()}").strip()
    normalized = unicodedata.normalize("NFKD", raw).encode("ascii", "ignore").decode("ascii")
    slug = re.sub(r"[^A-Za-z0-9_-]+", "-", normalized).strip("-").lower()
    return f"research-chat-{(slug[:80] or 'session')}.docx"


def _content_disposition_for_docx(title: str | None) -> str:
    original = (title or f"research-chat-{datetime.now(timezone.utc).date().isoformat()}").strip()
    if not original.lower().endswith(".docx"):
        original = f"{original}.docx"
    return (
        f'attachment; filename="{_ascii_export_filename(title)}"; '
        f"filename*=UTF-8''{quote(original.encode('utf-8'))}"
    )


@router.get("/research-sessions/{session_id}/export.docx")
async def export_research_session_docx(session_id: str, user: dict = Depends(get_current_user)):
    from io import BytesIO

    from docx import Document
    from fastapi.responses import StreamingResponse

    user_id = _get_user_id(user)
    session = _get_owned_session(session_id, user_id)
    try:
        resp = (
            supabase.table("research_session_messages")
            .select("id, role, content, citations, created_at")
            .eq("research_session_id", session_id)
            .order("created_at", desc=False)
            .execute()
        )
    except Exception as exc:
        logger.exception("Export research session messages failed")
        raise HTTPException(status_code=500, detail={"code": "INTERNAL_ERROR", "message": "Không thể tạo file chia sẻ."}) from exc
    rows, error = _supabase_response_data(resp)
    if error:
        raise HTTPException(status_code=500, detail={"code": "INTERNAL_ERROR", "message": "Không thể tạo file chia sẻ."})

    doc = Document()
    doc.add_heading(session.get("title") or "Lịch sử nghiên cứu", 0)
    doc.add_paragraph(f"Ngày xuất file: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}")

    for row in rows or []:
        role = "User" if row.get("role") == "user" else "Assistant"
        p = doc.add_paragraph()
        run = p.add_run(role)
        run.bold = True
        doc.add_paragraph(row.get("content") or "")
        citations = row.get("citations") or []
        if row.get("role") == "assistant" and citations:
            source_title = doc.add_paragraph("Nguồn tham khảo")
            source_title.runs[0].bold = True
            for index, citation in enumerate(citations, start=1):
                label = citation.get("citation_index") or index
                title = citation.get("document_title") or citation.get("filename") or "Tài liệu"
                page_start = citation.get("page_start") or citation.get("page")
                page_end = citation.get("page_end") or page_start
                score = citation.get("score")
                page_text = f", tr. {page_start}-{page_end}" if page_start and page_end and page_end != page_start else (f", tr. {page_start}" if page_start else "")
                score_text = f", score {score}" if score is not None else ""
                doc.add_paragraph(f"[{label}] {title}{page_text}{score_text}", style="List Bullet")
        doc.add_paragraph("")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": _content_disposition_for_docx(session.get("title"))},
    )


@router.post("/research-sessions/{session_id}/flashcards/generate", response_model=dict)
async def generate_research_session_flashcards(
    session_id: str,
    body: GenerateFlashcardsRequest,
    user: dict = Depends(get_current_user),
):
    from app.services.groq_service import generate_flashcards_from_context

    user_id = _get_user_id(user)
    session = _get_owned_session(session_id, user_id)
    selected = _selected_session_documents(session, body.selected_document_ids)
    rag_query = "Tạo flashcards ôn tập từ các khái niệm, luận điểm, phương pháp và kết quả quan trọng trong tài liệu đã chọn."
    context, warning = await _load_generation_context(session, selected, rag_query)
    try:
        flashcards = await generate_flashcards_from_context(context, body.count)
    except Exception as exc:
        message = str(exc) or "Thiếu GROQ_API_KEY hoặc không thể tạo flashcards."
        raise HTTPException(status_code=500, detail={"code": "GROQ_FAILED", "message": message}) from exc

    return {"success": True, "data": {"flashcards": flashcards, "warning": warning}}


@router.post("/research-sessions/{session_id}/quizzes/generate", response_model=dict)
async def generate_research_session_quiz(
    session_id: str,
    body: GenerateQuizRequest,
    user: dict = Depends(get_current_user),
):
    from app.services.groq_service import generate_quiz_from_context

    if body.question_type not in {"mixed", "multiple_choice", "true_false"}:
        raise HTTPException(status_code=400, detail={"code": "INVALID_QUESTION_TYPE", "message": "question_type phải là mixed, multiple_choice hoặc true_false."})
    user_id = _get_user_id(user)
    session = _get_owned_session(session_id, user_id)
    selected = _selected_session_documents(session, body.selected_document_ids)
    rag_query = "Tạo câu hỏi trắc nghiệm ôn tập từ các khái niệm, luận điểm, phương pháp và kết quả quan trọng trong tài liệu đã chọn."
    context, warning = await _load_generation_context(session, selected, rag_query)
    try:
        questions = await generate_quiz_from_context(context, body.count, body.question_type)
    except Exception as exc:
        message = str(exc) or "Thiếu GROQ_API_KEY hoặc không thể tạo quiz/test."
        raise HTTPException(status_code=502, detail={"code": "GROQ_FAILED", "message": message}) from exc
    return {"success": True, "data": {"quiz": {"id": f"quiz-{session_id}", "title": "Bộ câu hỏi trắc nghiệm", "questions": questions}, "questions": questions, "warning": warning}}


@router.post("/research-sessions/{session_id}/tests/generate", response_model=dict)
async def generate_research_session_test(
    session_id: str,
    body: GenerateTestRequest,
    user: dict = Depends(get_current_user),
):
    from app.services.groq_service import generate_test_from_context

    if body.count != 10:
        raise HTTPException(status_code=400, detail={"code": "INVALID_TEST_COUNT", "message": "Tạo bài kiểm tra yêu cầu đúng 10 câu hỏi."})
    user_id = _get_user_id(user)
    session = _get_owned_session(session_id, user_id)
    selected = _selected_session_documents(session, body.selected_document_ids)
    rag_query = "Tạo bài kiểm tra 10 câu phối hợp multiple choice, true false, điền từ và tự luận từ tài liệu đã chọn."
    context, warning = await _load_generation_context(session, selected, rag_query)
    try:
        test = await generate_test_from_context(context, 10)
    except Exception as exc:
        message = str(exc) or "Thiếu GROQ_API_KEY hoặc không thể tạo quiz/test."
        raise HTTPException(status_code=502, detail={"code": "GROQ_FAILED", "message": message}) from exc
    return {"success": True, "data": {"test": test, "warning": warning}}
